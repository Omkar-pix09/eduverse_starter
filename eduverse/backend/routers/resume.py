import os
import re
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from pydantic import BaseModel
from typing import List, Optional

import models
from database import get_db

router = APIRouter(prefix="/api/resume", tags=["resume"])

GENERATED_DIR = os.path.join(os.path.dirname(__file__), "..", "uploads", "resumes")
os.makedirs(GENERATED_DIR, exist_ok=True)


class ResumeBuildRequest(BaseModel):
    target_role: Optional[str] = None
    job_description: Optional[str] = None  # used for ATS matching


def _tokenize(text: str) -> set:
    return set(re.findall(r"[a-zA-Z0-9\+\#\.]+", text.lower()))


@router.post("/{user_id}/build")
def build_resume(user_id: int, req: ResumeBuildRequest, db: Session = Depends(get_db)):
    profile = db.query(models.Profile).filter(models.Profile.user_id == user_id).first()
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not profile or not user:
        raise HTTPException(status_code=404, detail="Profile not found")

    skills = profile.skills or []
    certs = profile.certifications or []

    # --- Generate a clean .docx resume ---
    doc = Document()
    doc.add_heading(user.full_name, level=1)
    doc.add_paragraph(f"{user.email}")
    if req.target_role:
        doc.add_paragraph(f"Target Role: {req.target_role}")

    doc.add_heading("Education", level=2)
    doc.add_paragraph(f"Branch: {profile.branch or 'N/A'}  |  CGPA: {profile.cgpa or 'N/A'}  |  Grad Year: {profile.grad_year or 'N/A'}")

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(skills) if skills else "N/A")

    doc.add_heading("Certifications", level=2)
    doc.add_paragraph(", ".join(certs) if certs else "N/A")

    if profile.github_username:
        doc.add_heading("Projects (from GitHub)", level=2)
        github_data = (
            db.query(models.GitHubData)
            .filter(models.GitHubData.user_id == user_id)
            .order_by(models.GitHubData.synced_at.desc())
            .first()
        )
        if github_data and github_data.top_repos:
            for repo in github_data.top_repos:
                doc.add_paragraph(f"{repo['name']} — {repo['language'] or 'N/A'} ({repo['stars']} stars)", style="List Bullet")

    filename = f"resume_user{user_id}.docx"
    filepath = os.path.join(GENERATED_DIR, filename)
    doc.save(filepath)

    # --- ATS score against job description (if provided) ---
    ats_score, matched, missing = None, [], []
    if req.job_description:
        resume_text = " ".join(skills + certs + [profile.branch or "", req.target_role or ""])
        resume_tokens = _tokenize(resume_text)
        jd_tokens = _tokenize(req.job_description)

        matched = sorted(resume_tokens & jd_tokens)
        missing = sorted(jd_tokens - resume_tokens)

        vectorizer = TfidfVectorizer()
        tfidf = vectorizer.fit_transform([resume_text, req.job_description])
        ats_score = round(float(cosine_similarity(tfidf[0], tfidf[1])[0][0]) * 100, 2)

    record = models.Resume(
        user_id=user_id,
        target_role=req.target_role,
        content_json={"skills": skills, "certifications": certs},
        ats_score=ats_score,
        matched_keywords=matched,
        missing_keywords=missing,
    )
    db.add(record)
    db.commit()

    return {
        "resume_file": filename,
        "ats_score": ats_score,
        "matched_keywords": matched,
        "missing_keywords": missing,
    }


@router.get("/{user_id}/download/{filename}")
def download_resume(user_id: int, filename: str):
    filepath = os.path.join(GENERATED_DIR, filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="Resume not found")
    return FileResponse(filepath, filename=filename)
